from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from inventario_faces import __version__
from inventario_faces.domain.config import AppConfig
from inventario_faces.domain.entities import FileRecord, FaceCluster, FaceOccurrence, InventoryResult, ReportArtifacts
from inventario_faces.utils.latex import format_seconds
from inventario_faces.utils.path_utils import ensure_directory


PROJECT_URL = "https://github.com/demusis/inventario_faces"


class DocxReportGenerator:
    def __init__(self, config: AppConfig) -> None:
        self._config = config

    def generate(self, result: InventoryResult) -> ReportArtifacts:
        report_directory = ensure_directory(result.run_directory / "report")
        docx_path = report_directory / "relatorio_forense.docx"
        document = self._build_document(result)
        document.save(docx_path)
        return ReportArtifacts(tex_path=report_directory / "relatorio_forense.tex", pdf_path=None, docx_path=docx_path)

    def _build_document(self, result: InventoryResult) -> Document:
        document = Document()
        self._configure_styles(document)
        self._add_cover(document, result)
        document.add_page_break()
        self._add_heading(document, "Resumo Executivo", level=1)
        self._add_paragraph(document, self._executive_summary_text(result))
        self._add_notice_box(
            document,
            "Advertência pericial",
            (
                "Os resultados são probabilísticos e não constituem identificação conclusiva de indivíduos. "
                "Qualquer inferência deve ser submetida à revisão humana especializada e correlacionada com outros elementos de prova."
            ),
        )
        self._add_notice_box(document, "Cadeia de custódia", self._config.forensics.chain_of_custody_note)

        self._add_heading(document, "Metodologia", level=1)
        for step in self._methodology_items(result):
            self._add_list_item(document, step)

        self._add_heading(document, "Resultados", level=1)
        self._add_heading(document, "Estatísticas de tamanho das faces", level=2)
        self._add_face_statistics_table(document, result)
        self._add_heading(document, "Possíveis indivíduos", level=2)
        self._add_cluster_overview_table(document, result.clusters)
        for cluster in result.clusters:
            occurrences = [item for item in result.occurrences if item.cluster_id == cluster.cluster_id]
            self._add_cluster_section(document, cluster, occurrences)

        self._add_heading(document, "Anexo técnico", level=1)
        self._add_heading(document, "Hashes SHA-512", level=2)
        self._add_hashes_table(document, result.files)
        self._add_heading(document, "Características de imagens e vídeos (MediaInfo)", level=2)
        self._add_media_info_table(document, result.files)
        self._add_heading(document, "Parâmetros do modelo e da execução", level=2)
        for item in self._technical_parameter_items(result):
            self._add_list_item(document, item)
        self._add_heading(document, "Registros de execução", level=2)
        self._add_log_excerpt(document, result.logs_directory / "run.log")

        return document

    def _configure_styles(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Cambria"
        normal_style.font.size = Pt(11)
        heading1 = document.styles["Heading 1"]
        heading1.font.name = "Cambria"
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading2 = document.styles["Heading 2"]
        heading2.font.name = "Cambria"
        heading2.font.size = Pt(13)
        heading2.font.bold = True

    def _add_cover(self, document: Document, result: InventoryResult) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self._config.app.report_title)
        run.bold = True
        run.font.size = Pt(24)

        organization = document.add_paragraph()
        organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
        organization.add_run(self._config.app.organization).font.size = Pt(16)

        generated = document.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated.add_run(f"Gerado em {result.finished_at_utc.strftime('%Y-%m-%d %H:%M:%SZ')}").font.size = Pt(12)

        document.add_paragraph()
        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.add_run("Relatório técnico para apoio investigativo assistido por IA").italic = True

        caveat = document.add_paragraph()
        caveat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caveat.add_run("Este documento não deve ser interpretado como identificação conclusiva de indivíduos.")

    def _executive_summary_text(self, result: InventoryResult) -> str:
        summary = result.summary
        return (
            f"O presente relatório registra o processamento automatizado do diretório {result.root_directory}. "
            f"Foram catalogados e submetidos a cálculo de hash {summary.total_files} arquivo(s). No conjunto "
            f"catalogado, há {summary.media_files} mídia(s) suportada(s) para análise facial automatizada, com "
            f"a seguinte composição: {summary.image_files} imagem(ns) e {summary.video_files} vídeo(s). "
            f"A etapa de detecção registrou {summary.total_occurrences} ocorrência(s) facial(is). A etapa de "
            f"agrupamento resultou em {summary.total_clusters} possível(is) indivíduo(s). Foram assinalados "
            f"{summary.probable_match_pairs} par(es) como possivelmente correlato(s), sem caráter conclusivo de identificação."
        )

    def _methodology_items(self, result: InventoryResult) -> list[str]:
        detection_size = self._detection_size_label()
        max_frames_text = (
            str(self._config.video.max_frames_per_video)
            if self._config.video.max_frames_per_video is not None
            else "sem limite"
        )
        return [
            "Varredura recursiva do diretório de entrada, com registro individual de caminho, tamanho e hash SHA-512 para cada arquivo encontrado.",
            "Classificação de mídia em imagem, vídeo ou outro formato, sem alteração dos arquivos originais.",
            f"Para vídeos, amostragem temporal a cada {self._config.video.sampling_interval_seconds:.2f} segundos, limitada a {max_frames_text} quadros por arquivo.",
            f"Detecção facial e extração de vetores de características normalizados com mecanismo configurado em {self._config.face_model.backend} / {self._config.face_model.model_name}, com tamanho de detecção definido em {detection_size}.",
            f"Seleção de faces condicionada a qualidade mínima de {self._config.face_model.minimum_face_quality:.2f}, aferida pela pontuação de detecção do mecanismo facial.",
            f"Seleção complementar por tamanho mínimo da face, exigindo ao menos {self._config.face_model.minimum_face_size_pixels} pixels na menor dimensão da caixa delimitadora.",
            f"Agrupamento incremental por similaridade de cosseno, com limiar de atribuição {self._config.clustering.assignment_similarity:.2f} e limiar de sugestão entre possíveis indivíduos {self._config.clustering.candidate_similarity:.2f}.",
            f"Consolidação de vestígios forenses, registros e artefatos derivados em diretório dedicado de execução: {result.run_directory}.",
        ]

    def _technical_parameter_items(self, result: InventoryResult) -> list[str]:
        max_frames_text = (
            str(self._config.video.max_frames_per_video)
            if self._config.video.max_frames_per_video is not None
            else "sem limite"
        )
        providers = ", ".join(self._config.face_model.providers) or "automatico"
        mediainfo_directory = (
            self._config.app.mediainfo_directory or "nao configurado; resolucao automatica pelo PATH do sistema"
        )
        return [
            f"Aplicação: versão={__version__}; diretório de saída={self._config.app.output_directory_name}; nível de log={self._config.app.log_level}.",
            f"Mídias: extensões de imagem={', '.join(self._config.media.image_extensions)}; extensões de vídeo={', '.join(self._config.media.video_extensions)}.",
            f"Vídeo: intervalo de amostragem={self._config.video.sampling_interval_seconds:.2f} s; máximo de quadros por arquivo={max_frames_text}.",
            (
                f"Análise facial: mecanismo={self._config.face_model.backend}; modelo={self._config.face_model.model_name}; "
                f"tamanho de detecção={self._detection_size_label()}; qualidade mínima={self._config.face_model.minimum_face_quality:.2f}; "
                f"tamanho mínimo da face={self._config.face_model.minimum_face_size_pixels} px; contexto={self._config.face_model.ctx_id}; "
                f"mecanismos de execução={providers}."
            ),
            (
                f"Agrupamento: limiar de atribuição={self._config.clustering.assignment_similarity:.2f}; "
                f"limiar de sugestão entre possíveis indivíduos={self._config.clustering.candidate_similarity:.2f}; "
                f"tamanho mínimo do grupo={self._config.clustering.min_cluster_size}."
            ),
            (
                f"Relatório: faces máximas por possível indivíduo na galeria={self._config.reporting.max_gallery_faces_per_group}; "
                f"compilação automática do PDF={'sim' if self._config.reporting.compile_pdf else 'nao'}."
            ),
            f"MediaInfo: {self._mediainfo_status(result.files)}; diretório configurado={mediainfo_directory}.",
            f"Software utilizado: aplicação de código aberto disponível em {PROJECT_URL}.",
        ]

    def _add_heading(self, document: Document, text: str, level: int) -> None:
        document.add_heading(text, level=level)

    def _add_paragraph(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(8)

    def _add_list_item(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(text)

    def _add_notice_box(self, document: Document, title: str, body: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        title_run = paragraph.add_run(f"{title}. ")
        title_run.bold = True
        paragraph.add_run(body)
        document.add_paragraph()

    def _add_face_statistics_table(self, document: Document, result: InventoryResult) -> None:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = [
            "Conjunto",
            "Número de faces",
            "Mínimo (px)",
            "Máximo (px)",
            "Média (px)",
            "Desvio padrão (px)",
        ]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for label, stats in [
            ("Faces detectadas antes dos filtros", result.summary.total_detected_face_sizes),
            ("Faces selecionadas após os filtros", result.summary.selected_face_sizes),
        ]:
            row = table.add_row().cells
            row[0].text = label
            if stats.count == 0:
                row[1].text = "0"
                row[2].text = "-"
                row[3].text = "-"
                row[4].text = "-"
                row[5].text = "-"
            else:
                row[1].text = str(stats.count)
                row[2].text = f"{stats.min_pixels:.1f}"
                row[3].text = f"{stats.max_pixels:.1f}"
                row[4].text = f"{stats.mean_pixels:.1f}"
                row[5].text = f"{stats.stddev_pixels:.1f}"
        document.add_paragraph()

    def _add_cluster_overview_table(self, document: Document, clusters: list[FaceCluster]) -> None:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ["Identificador", "Ocorrências", "Possíveis indivíduos correlatos"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        if not clusters:
            row = table.add_row().cells
            row[0].text = "-"
            row[1].text = "0"
            row[2].text = "Nenhum possível indivíduo identificado."
        for cluster in clusters:
            row = table.add_row().cells
            row[0].text = cluster.cluster_id
            row[1].text = str(len(cluster.occurrence_ids))
            row[2].text = ", ".join(cluster.candidate_cluster_ids) or "-"
        document.add_paragraph()

    def _add_cluster_section(self, document: Document, cluster: FaceCluster, occurrences: list[FaceOccurrence]) -> None:
        self._add_heading(document, f"Possível indivíduo {cluster.cluster_id}", level=3)
        self._add_paragraph(document, f"Total de ocorrências: {len(occurrences)}")
        self._add_paragraph(
            document,
            f"Possíveis indivíduos correlatos: {', '.join(cluster.candidate_cluster_ids) or '-'}",
        )
        self._add_heading(document, "Galeria comparativa", level=4)
        if occurrences:
            for occurrence in occurrences[: self._config.reporting.max_gallery_faces_per_group]:
                self._add_gallery_row(document, occurrence)
        else:
            self._add_paragraph(document, "Nenhuma ilustração disponível para este possível indivíduo.")
        self._add_heading(document, "Ocorrências do possível indivíduo", level=4)
        self._add_occurrences_table(document, occurrences)

    def _add_gallery_row(self, document: Document, occurrence: FaceOccurrence) -> None:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        widths = [Inches(1.7), Inches(3.6), Inches(2.4)]
        labels = ["Recorte facial", "Imagem ou quadro de origem", "Dados da ocorrência"]
        header = table.add_row().cells
        for index, label in enumerate(labels):
            header[index].text = label
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            cell.width = width
        self._fill_image_cell(cells[0], occurrence.crop_path, width=Inches(1.5))
        self._fill_image_cell(cells[1], occurrence.context_image_path, width=Inches(3.2))
        self._fill_metadata_cell(cells[2], occurrence)
        document.add_paragraph()

    def _fill_image_cell(self, cell, image_path: Path | None, width: Inches) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if image_path is None or not image_path.exists():
            paragraph.add_run("Artefato não disponível")
            return
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=width)

    def _fill_metadata_cell(self, cell, occurrence: FaceOccurrence) -> None:
        lines = [
            occurrence.occurrence_id,
            f"Arquivo: {occurrence.source_path.name}",
            f"Marca temporal: {format_seconds(occurrence.frame_timestamp_seconds)}",
            f"Pontuação: {occurrence.detection_score:.3f}",
            (
                f"Caixa delimitadora: {occurrence.bbox.x1:.1f}, {occurrence.bbox.y1:.1f}, "
                f"{occurrence.bbox.x2:.1f}, {occurrence.bbox.y2:.1f}"
            ),
        ]
        cell.text = ""
        for index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.add_run(line)

    def _add_occurrences_table(self, document: Document, occurrences: list[FaceOccurrence]) -> None:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Ocorrência", "Arquivo", "Marca temporal", "Pontuação"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        if not occurrences:
            row = table.add_row().cells
            row[0].text = "-"
            row[1].text = "Nenhuma ocorrência detalhada."
            row[2].text = "-"
            row[3].text = "-"
            return
        for item in occurrences:
            row = table.add_row().cells
            row[0].text = item.occurrence_id
            row[1].text = item.source_path.name
            row[2].text = format_seconds(item.frame_timestamp_seconds)
            row[3].text = f"{item.detection_score:.3f}"
        document.add_paragraph()

    def _add_hashes_table(self, document: Document, files: list[FileRecord]) -> None:
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Arquivo"
        table.rows[0].cells[1].text = "SHA-512"
        if not files:
            row = table.add_row().cells
            row[0].text = "-"
            row[1].text = "Nenhum arquivo catalogado."
            return
        for item in files:
            row = table.add_row().cells
            row[0].text = str(item.path)
            row[1].text = item.sha512
        document.add_paragraph()

    def _add_media_info_table(self, document: Document, files: list[FileRecord]) -> None:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ["Arquivo", "Fluxo", "Características"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        media_files = [item for item in files if item.media_type.value in {"image", "video"}]
        if not media_files:
            row = table.add_row().cells
            row[0].text = "-"
            row[1].text = "-"
            row[2].text = "Nenhuma mídia elegível para extração via MediaInfo."
            document.add_paragraph()
            return
        for item in media_files:
            if item.media_info_tracks:
                for index, track in enumerate(item.media_info_tracks):
                    row = table.add_row().cells
                    row[0].text = str(item.path) if index == 0 else ""
                    row[1].text = track.track_type
                    row[2].text = "; ".join(f"{attribute.label}: {attribute.value}" for attribute in track.attributes)
            else:
                row = table.add_row().cells
                row[0].text = str(item.path)
                row[1].text = "-"
                row[2].text = item.media_info_error or "MediaInfo não disponível para este arquivo."
        document.add_paragraph()

    def _add_log_excerpt(self, document: Document, log_path: Path, max_lines: int = 60) -> None:
        if not log_path.exists():
            self._add_paragraph(document, "Registro principal não encontrado.")
            return
        lines = log_path.read_text(encoding="utf-8", errors="ignore").splitlines()[-max_lines:]
        if not lines:
            self._add_paragraph(document, "Nenhum registro disponível.")
            return
        paragraph = document.add_paragraph()
        for index, line in enumerate(lines):
            paragraph.add_run(line)
            if index != len(lines) - 1:
                paragraph.add_run("\n")

    def _detection_size_label(self) -> str:
        if self._config.face_model.det_size is None:
            return "resolução original do arquivo ou quadro"
        return f"{self._config.face_model.det_size[0]} x {self._config.face_model.det_size[1]}"

    def _mediainfo_status(self, files: list[FileRecord]) -> str:
        media_files = [item for item in files if item.media_type.value in {"image", "video"}]
        if not media_files:
            return "nenhuma mídia elegível para coleta"
        if any(item.media_info_tracks for item in media_files):
            return "coleta executada para mídias elegíveis, com resultados detalhados na subseção anterior"
        error_messages = [item.media_info_error for item in media_files if item.media_info_error]
        if error_messages:
            return f"coleta não realizada ou indisponível no ambiente; motivo predominante={error_messages[0]}"
        return "coleta não realizada"
