from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from inventario_faces.domain.config import AppConfig
from inventario_faces.domain.entities import (
    FaceSearchMatch,
    FaceSearchQuery,
    FaceSearchQueryEvent,
    FaceSearchResult,
    ReportArtifacts,
)
from inventario_faces.reporting.report_support import (
    face_search_methodology_items,
    software_reference_abnt_text,
    technical_parameter_items,
)
from inventario_faces.utils.latex import format_seconds
from inventario_faces.utils.path_utils import ensure_directory
from inventario_faces.utils.time_utils import format_local_datetime


class FaceSearchDocxReportGenerator:
    def __init__(self, config: AppConfig) -> None:
        self._config = config

    def generate(self, result: FaceSearchResult) -> ReportArtifacts:
        report_directory = ensure_directory(result.inventory_result.run_directory / "report")
        docx_path = report_directory / "relatorio_busca_por_face.docx"
        document = self._build_document(result)
        document.save(docx_path)
        return ReportArtifacts(
            tex_path=report_directory / "relatorio_busca_por_face.tex",
            pdf_path=None,
            docx_path=docx_path,
        )

    def _build_document(self, result: FaceSearchResult) -> Document:
        document = Document()
        document.styles["Normal"].font.name = "Cambria"
        document.styles["Normal"].font.size = Pt(11)

        self._add_report_header(
            document,
            f"{self._config.app.report_title} - Busca por Faces",
            result.inventory_result.finished_at_utc,
        )

        self._add_heading(document, "Resumo Executivo", 1)
        queries = self._queries(result)
        if not queries:
            summary_text = (
                f"A busca facial recebeu {result.summary.query_image_count} arquivo(s) de consulta. "
                "Nenhum deles gerou face de referência utilizável para a etapa de busca vetorial. "
                "Todos os eventos das consultas, inclusive erros de leitura, corrupção e descartes por critérios técnicos, "
                "foram individualizados na seção de arquivos de consulta informados. "
            )
            search_outcome = (
                f"O diretório {result.inventory_result.root_directory} foi processado, mas a busca vetorial "
                "não foi executada por ausência de referência facial válida."
            )
        elif len(queries) == 1 and result.summary.query_image_count == 1:
            summary_text = (
                f"A consulta utilizou o arquivo {queries[0].source_path}. "
                f"Foram detectadas {result.summary.query_faces_detected} face(s) elegíveis na imagem de consulta "
                f"e a face selecionada automaticamente para a pesquisa corresponde ao track {queries[0].selected_track_id}. "
            )
            search_outcome = (
                f"A varredura retornou {result.summary.compatible_tracks} track(s) compatíveis, distribuídos em "
                f"{result.summary.compatible_clusters} grupo(s), com limiar probabilístico de "
                f"{result.summary.compatibility_threshold:.2f}."
            )
        else:
            summary_text = (
                f"A busca facial recebeu {result.summary.query_image_count} arquivo(s) de consulta. "
                f"{result.summary.query_faces_selected} arquivo(s) geraram face(s) de referência utilizável(is), com "
                f"{result.summary.query_faces_detected} face(s) elegíveis detectadas nas consultas válidas. "
                f"{result.summary.query_images_rejected} arquivo(s) foram descartados, permanecendo individualizados "
                "com o motivo do descarte na seção de arquivos de consulta informados. "
            )
            search_outcome = (
                f"A varredura retornou {result.summary.compatible_tracks} track(s) compatíveis, distribuídos em "
                f"{result.summary.compatible_clusters} grupo(s), com limiar probabilístico de "
                f"{result.summary.compatibility_threshold:.2f}."
            )
        self._add_paragraph(
            document,
            f"{summary_text}{search_outcome}",
        )
        self._add_notice(
            document,
            "Advertência pericial",
            "A busca abaixo aponta apenas compatibilidade probabilística e não constitui prova conclusiva de identidade.",
        )

        self._add_heading(document, "Metodologia da Busca", 1)
        for item in face_search_methodology_items(
            result.summary.query_image_count,
            [query.selected_track_id for query in queries],
        ):
            self._add_list_item(document, item)

        self._add_heading(document, "Arquivos de Consulta Informados", 1)
        self._add_query_table(document, result)

        self._add_heading(document, "Resultados Compatíveis", 1)
        self._add_matches_table(document, result.matches)

        self._add_heading(document, "Anexo técnico", 1)
        self._add_heading(document, "Parâmetros e melhorias aplicadas", 2)
        for item in technical_parameter_items(self._config, result.inventory_result.search):
            self._add_list_item(document, item)
        self._add_list_item(document, f"Referência do software: {software_reference_abnt_text()}")
        self._add_heading(document, "Rastreabilidade da busca", 2)
        report_path = result.inventory_result.report.pdf_path or result.inventory_result.report.docx_path or result.inventory_result.report.tex_path
        self._add_paragraph(document, f"Inventário de referência usado na busca: {report_path}")
        self._add_paragraph(document, f"Exportação estruturada da busca: {result.export_path or '-'}")
        return document

    def _add_report_header(self, document: Document, title_text: str, finished_at_utc) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(16)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"{self._config.app.organization} | Emitido em {format_local_datetime(finished_at_utc)}"
        )
        meta_run.font.size = Pt(10)

    def _add_query_table(self, document: Document, result: FaceSearchResult) -> None:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_cell_text(table.rows[0].cells[0], "Consulta", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        self._set_cell_text(table.rows[0].cells[1], "Situação", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        self._set_cell_text(table.rows[0].cells[2], "Recorte da consulta", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        self._set_cell_text(table.rows[0].cells[3], "Quadro de origem da consulta", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        self._set_cell_text(table.rows[0].cells[4], "Metadados da consulta", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

        for event in self._query_events(result):
            row = table.add_row().cells
            self._set_cell_text(row[0], str(event.query_index), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(
                row[1],
                "Selecionada" if event.status == "selected" else "Descartada",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._add_cell_image(row[2], event.crop_path, width=Inches(1.35))
            self._add_cell_image(row[3], event.context_image_path, width=Inches(2.1))
            quality = "-" if event.quality_score is None else f"{event.quality_score:.3f}"
            detected_faces = "-" if event.detected_face_count is None else str(event.detected_face_count)
            self._set_cell_text(
                row[4],
                "\n".join(
                    [
                        f"Arquivo consultado: {event.source_path}",
                        f"SHA-512: {event.sha512 or '-'}",
                        f"Situação da consulta: {'Selecionada' if event.status == 'selected' else 'Descartada'}",
                        f"Faces elegíveis detectadas: {detected_faces}",
                        f"Track selecionado: {event.selected_track_id or '-'}",
                        f"Ocorrência selecionada: {event.selected_occurrence_id or '-'}",
                        f"Keyframe de referência: {event.selected_keyframe_id or '-'}",
                        f"Qualidade da face consultada: {quality}",
                        (
                            f"Evento reportado: {event.error_type} - {event.error_message}"
                            if event.error_message is not None and event.error_type is not None
                            else (
                                f"Evento reportado: {event.error_message}"
                                if event.error_message is not None
                                else "Evento reportado: processamento concluído sem erro."
                            )
                        ),
                    ]
                ),
            )
        document.add_paragraph()

    def _add_matches_table(self, document: Document, matches: list[FaceSearchMatch]) -> None:
        if not matches:
            self._add_paragraph(document, "Nenhum track compatível superou o limiar probabilístico configurado para as consultas.")
            return

        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = table.rows[0].cells
        for index, label in enumerate(("Pos.", "Grupo", "Track", "Recorte", "Quadro de origem", "Metadados")):
            self._set_cell_text(headers[index], label, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

        for match in matches:
            row = table.add_row().cells
            self._set_cell_text(row[0], str(match.rank), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row[1], match.cluster_id or "-", alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_text(row[2], match.track_id, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._add_cell_image(row[3], match.crop_path, width=Inches(1.2))
            self._add_cell_image(row[4], match.context_image_path, width=Inches(2.1))
            self._set_cell_text(
                row[5],
                "\n".join(self._match_metadata_lines(match)),
            )
        document.add_paragraph()

    def _match_metadata_lines(self, match: FaceSearchMatch) -> list[str]:
        occurrence_score = "-" if match.occurrence_score is None else f"{match.occurrence_score:.3f}"
        cluster_score = "-" if match.cluster_score is None else f"{match.cluster_score:.3f}"
        return [
            f"Consulta que sustentou o score: {match.query_source_path}" if match.query_source_path is not None else "Consulta que sustentou o score: -",
            f"Track da consulta: {match.query_selected_track_id or '-'}",
            f"Ocorrência da consulta: {match.query_selected_occurrence_id or '-'}",
            f"Origem: {match.source_path.name}",
            f"Ocorrência: {match.occurrence_id or '-'}",
            f"Instante da ocorrência: {format_seconds(match.timestamp_seconds)}",
            f"Intervalo do track: {format_seconds(match.track_start_time)} - {format_seconds(match.track_end_time)}",
            f"Quadro da ocorrência: {match.frame_index:06d}" if match.frame_index is not None else "Quadro da ocorrência: -",
            f"Similaridade do grupo: {cluster_score}",
            f"Similaridade do track: {match.track_score:.3f}",
            f"Similaridade da ocorrência: {occurrence_score}",
        ]

    def _queries(self, result: FaceSearchResult) -> list[FaceSearchQuery]:
        if result.queries:
            return list(result.queries)
        return [result.query] if result.query is not None else []

    def _query_events(self, result: FaceSearchResult) -> list[FaceSearchQueryEvent]:
        if result.query_events:
            return list(result.query_events)
        return [
            FaceSearchQueryEvent(
                query_index=query.query_index,
                source_path=query.source_path,
                status="selected",
                sha512=query.sha512,
                detected_face_count=query.detected_face_count,
                selected_track_id=query.selected_track_id,
                selected_occurrence_id=query.selected_occurrence_id,
                selected_keyframe_id=query.selected_keyframe_id,
                crop_path=query.crop_path,
                context_image_path=query.context_image_path,
                quality_score=query.quality_score,
            )
            for query in self._queries(result)
        ]

    def _add_heading(self, document: Document, text: str, level: int) -> None:
        document.add_heading(text, level=level)

    def _add_paragraph(self, document: Document, text: str) -> None:
        document.add_paragraph(text)

    def _add_list_item(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(text)

    def _add_notice(self, document: Document, title: str, body: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        title_run = paragraph.add_run(f"{title}. ")
        title_run.bold = True
        paragraph.add_run(body)
        document.add_paragraph()

    def _set_cell_text(
        self,
        cell,
        text: str,
        *,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        bold: bool = False,
    ) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.bold = bold

    def _add_cell_image(self, cell, image_path: Path | None, *, width: Inches) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if image_path is None or not image_path.exists():
            paragraph.add_run("Artefato não disponível")
            return
        paragraph.add_run().add_picture(str(image_path), width=width)
