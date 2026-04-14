from __future__ import annotations

import tempfile
import unittest
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from tests._bootstrap import PROJECT_ROOT  # noqa: F401
from inventario_faces.domain.config import (
    AppConfig,
    AppSettings,
    ClusteringSettings,
    FaceModelSettings,
    ForensicsSettings,
    MediaSettings,
    ReportingSettings,
    VideoSettings,
)
from inventario_faces.domain.entities import (
    BoundingBox,
    FaceCluster,
    FaceOccurrence,
    FaceSearchQuery,
    FaceSearchQueryEvent,
    FaceSearchResult,
    FaceSearchSummary,
    FaceSizeStatistics,
    FaceTrack,
    FileRecord,
    InventoryResult,
    KeyFrame,
    MediaType,
    ProcessingSummary,
    ReportArtifacts,
    SearchArtifacts,
    TrackQualityStatistics,
)
from inventario_faces.reporting.face_search_docx_renderer import FaceSearchDocxReportGenerator
from inventario_faces.reporting.face_search_latex_renderer import FaceSearchLatexReportGenerator


class _UnusedCompiler:
    def compile(self, tex_path: Path) -> Path:
        return tex_path.with_suffix(".pdf")


class FaceSearchReportGeneratorTests(unittest.TestCase):
    def test_docx_report_lists_rejected_query_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True, exist_ok=True)
            generator = FaceSearchDocxReportGenerator(self._config())

            artifacts = generator.generate(self._result(root, logs_dir))

            document = Document(str(artifacts.docx_path))
            full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            full_text += "\n" + "\n".join(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("Arquivos de Consulta Informados", full_text)
            self.assertIn("consulta_corrompida.jpg", full_text)
            self.assertIn("MediaDecodeError - Nao foi possivel ler a imagem", full_text)
            self.assertIn("Situação da consulta: Descartada", full_text)

    def test_latex_report_lists_rejected_query_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True, exist_ok=True)
            generator = FaceSearchLatexReportGenerator(self._config(), _UnusedCompiler())

            artifacts = generator.generate(self._result(root, logs_dir))

            tex_content = artifacts.tex_path.read_text(encoding="utf-8")
            self.assertIn("Arquivos de Consulta Informados", tex_content)
            self.assertIn("consulta\\_\\allowbreak{}corrompida.\\allowbreak{}jpg", tex_content)
            self.assertIn("MediaDecodeError - Nao foi possivel ler a imagem", tex_content)
            self.assertIn("Situa", tex_content)

    def _config(self) -> AppConfig:
        return AppConfig(
            app=AppSettings(
                name="Inventario Faces",
                output_directory_name="inventario_faces_output",
                report_title="Relatório de Busca Facial",
                organization="Laboratório Teste",
                log_level="INFO",
            ),
            media=MediaSettings(
                image_extensions=(".jpg", ".png"),
                video_extensions=(".mp4", ".avi"),
            ),
            video=VideoSettings(
                sampling_interval_seconds=1.0,
                max_frames_per_video=10,
            ),
            face_model=FaceModelSettings(
                backend="insightface",
                model_name="buffalo_l",
                det_size=(640, 640),
                minimum_face_quality=0.6,
                minimum_face_size_pixels=40,
                ctx_id=0,
                providers=("CPUExecutionProvider",),
            ),
            clustering=ClusteringSettings(
                assignment_similarity=0.52,
                candidate_similarity=0.44,
                min_cluster_size=1,
            ),
            reporting=ReportingSettings(
                compile_pdf=False,
            ),
            forensics=ForensicsSettings(
                chain_of_custody_note="Arquivos originais preservados."
            ),
        )

    def _result(self, root: Path, logs_dir: Path) -> FaceSearchResult:
        occurrence = FaceOccurrence(
            occurrence_id="O000001",
            source_path=Path("evidencia.mp4"),
            sha512="a" * 128,
            media_type=MediaType.VIDEO,
            analysis_timestamp_utc=datetime.now(UTC),
            frame_index=10,
            frame_timestamp_seconds=10.0,
            bbox=BoundingBox(10, 20, 60, 90),
            detection_score=0.83,
            embedding=[1.0, 0.0, 0.0],
            crop_path=None,
            context_image_path=None,
            cluster_id="I001",
            track_id="T000001",
        )
        cluster = FaceCluster(
            cluster_id="I001",
            track_ids=["T000001"],
            occurrence_ids=["O000001"],
            centroid_embedding=[1.0, 0.0, 0.0],
        )
        track = FaceTrack(
            track_id="T000001",
            source_path=Path("evidencia.mp4"),
            video_path=Path("evidencia.mp4"),
            media_type=MediaType.VIDEO,
            sha512="a" * 128,
            start_frame=10,
            end_frame=10,
            start_time=10.0,
            end_time=10.0,
            occurrence_ids=["O000001"],
            keyframe_ids=["K000001"],
            representative_embeddings=[[1.0, 0.0, 0.0]],
            average_embedding=[1.0, 0.0, 0.0],
            best_occurrence_id="O000001",
            quality_statistics=TrackQualityStatistics(
                total_detections=1,
                keyframe_count=1,
                mean_detection_score=0.83,
                max_detection_score=0.83,
                mean_quality_score=0.80,
                best_quality_score=0.80,
                mean_sharpness=0.70,
                mean_brightness=0.55,
                mean_illumination=0.90,
                mean_frontality=0.88,
                duration_seconds=0.0,
            ),
            cluster_id="I001",
        )
        keyframe = KeyFrame(
            keyframe_id="K000001",
            track_id="T000001",
            occurrence_id="O000001",
            source_path=Path("evidencia.mp4"),
            frame_index=10,
            timestamp_seconds=10.0,
            selection_reasons=("track_start",),
            detection_score=0.83,
            embedding=[1.0, 0.0, 0.0],
        )
        file_record = FileRecord(
            path=Path("evidencia.mp4"),
            media_type=MediaType.VIDEO,
            sha512="a" * 128,
            size_bytes=123456,
            discovered_at_utc=datetime.now(UTC),
            modified_at_utc=datetime.now(UTC),
        )
        summary = ProcessingSummary(
            total_files=1,
            media_files=1,
            image_files=0,
            video_files=1,
            total_occurrences=1,
            total_clusters=1,
            probable_match_pairs=0,
            total_tracks=1,
            total_keyframes=1,
            total_detected_face_sizes=FaceSizeStatistics(count=1, min_pixels=50, max_pixels=50, mean_pixels=50, stddev_pixels=0),
            selected_face_sizes=FaceSizeStatistics(count=1, min_pixels=50, max_pixels=50, mean_pixels=50, stddev_pixels=0),
        )
        inventory_result = InventoryResult(
            run_directory=root,
            started_at_utc=datetime.now(UTC),
            finished_at_utc=datetime.now(UTC),
            root_directory=root,
            files=[file_record],
            occurrences=[occurrence],
            clusters=[cluster],
            report=ReportArtifacts(tex_path=root / "report" / "relatorio_forense.tex", pdf_path=None),
            summary=summary,
            logs_directory=logs_dir,
            manifest_path=root / "inventory" / "manifest.json",
            tracks=[track],
            keyframes=[keyframe],
            search=SearchArtifacts(
                engine="numpy",
                track_index_path=None,
                track_metadata_path=None,
                cluster_index_path=None,
                cluster_metadata_path=None,
                track_vector_count=1,
                cluster_vector_count=1,
            ),
        )
        query = FaceSearchQuery(
            source_path=Path("consulta_valida.jpg"),
            sha512="b" * 128,
            detected_face_count=1,
            selected_track_id="Q001_T000001",
            selected_occurrence_id="Q001_O000001",
            selected_keyframe_id="Q001_K000001",
            crop_path=None,
            context_image_path=None,
            quality_score=0.77,
            query_index=1,
        )
        query_events = [
            FaceSearchQueryEvent(
                query_index=1,
                source_path=Path("consulta_valida.jpg"),
                status="selected",
                sha512="b" * 128,
                detected_face_count=1,
                selected_track_id="Q001_T000001",
                selected_occurrence_id="Q001_O000001",
                selected_keyframe_id="Q001_K000001",
                quality_score=0.77,
            ),
            FaceSearchQueryEvent(
                query_index=2,
                source_path=Path("consulta_corrompida.jpg"),
                status="rejected",
                error_type="MediaDecodeError",
                error_message="Nao foi possivel ler a imagem: consulta_corrompida.jpg",
            ),
        ]
        return FaceSearchResult(
            inventory_result=inventory_result,
            query=query,
            queries=[query],
            query_events=query_events,
            matches=[],
            summary=FaceSearchSummary(
                query_faces_detected=1,
                compatible_clusters=0,
                compatible_tracks=0,
                compatible_occurrences=0,
                compatibility_threshold=0.44,
                query_image_count=2,
                query_faces_selected=1,
                query_images_rejected=1,
            ),
            report=ReportArtifacts(
                tex_path=root / "report" / "relatorio_busca_por_face.tex",
                pdf_path=None,
                docx_path=root / "report" / "relatorio_busca_por_face.docx",
            ),
            export_path=root / "inventory" / "face_search.json",
        )


if __name__ == "__main__":
    unittest.main()
