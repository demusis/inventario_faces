from __future__ import annotations

import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document

from tests._bootstrap import PROJECT_ROOT  # noqa: F401
from tests._report_fixtures import build_inventory_result, build_report_config
from inventario_faces.infrastructure.latex_compiler import LatexCompiler
from inventario_faces.reporting.docx_renderer import DocxReportGenerator
from inventario_faces.reporting.latex_renderer import LatexReportGenerator


def _docx_full_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return text


class LatexRenderingTests(unittest.TestCase):
    def test_special_characters_are_escaped_in_tex(self) -> None:
        config = build_report_config(
            chain_of_custody_note="Nota com 100% de cadeia & custos #1 em R$_teste.",
        )
        generator = LatexReportGenerator(config, LatexCompiler())
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            result = build_inventory_result(
                root,
                logs_dir,
                source_name="video_50%_&_relatorio_#1.mp4",
            )

            artifacts = generator.generate(result)

            tex = artifacts.tex_path.read_text(encoding="utf-8")
            self.assertIn(r"100\% de cadeia \& custos \#1 em R\$\_teste", tex)
            self.assertNotIn("Nota com 100% de cadeia", tex)

    def test_tracks_above_limit_produce_truncation_notice(self) -> None:
        config = build_report_config(max_tracks_per_group=2)
        generator = LatexReportGenerator(config, LatexCompiler())
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            result = build_inventory_result(root, logs_dir, track_count=5)

            artifacts = generator.generate(result)

            tex = artifacts.tex_path.read_text(encoding="utf-8")
            self.assertIn("Exibindo 2 de 5 tracks deste grupo", tex)
            # Apenas os dois primeiros tracks entram na tabela.
            self.assertIn("T000002", tex)
            self.assertNotIn("T000003", tex)

    def test_tracks_within_limit_have_no_truncation_notice(self) -> None:
        config = build_report_config(max_tracks_per_group=8)
        generator = LatexReportGenerator(config, LatexCompiler())
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            result = build_inventory_result(root, logs_dir, track_count=3)

            artifacts = generator.generate(result)

            tex = artifacts.tex_path.read_text(encoding="utf-8")
            self.assertNotIn("Exibindo", tex)

    @unittest.skipIf(shutil.which("pdflatex") is None, "pdflatex indisponivel no PATH")
    def test_pdf_compilation_succeeds_with_pdflatex(self) -> None:
        config = build_report_config(compile_pdf=True)
        generator = LatexReportGenerator(config, LatexCompiler())
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            result = build_inventory_result(
                root,
                logs_dir,
                source_name="video_50%_&_relatorio_#1.mp4",
            )

            artifacts = generator.generate(result)

            self.assertIsNotNone(artifacts.pdf_path)
            self.assertTrue(artifacts.pdf_path.exists())
            self.assertGreater(artifacts.pdf_path.stat().st_size, 0)


class DocxRenderingTests(unittest.TestCase):
    def test_tracks_above_limit_produce_truncation_notice(self) -> None:
        config = build_report_config(max_tracks_per_group=2)
        generator = DocxReportGenerator(config)
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            (logs_dir / "run.log").write_text("linha de log", encoding="utf-8")
            result = build_inventory_result(root, logs_dir, track_count=5)

            artifacts = generator.generate(result)

            full_text = _docx_full_text(artifacts.docx_path)
            self.assertIn("Exibindo 2 de 5 tracks deste grupo", full_text)

    def test_special_characters_do_not_break_docx(self) -> None:
        config = build_report_config(
            chain_of_custody_note="Nota com 100% de cadeia & custos #1 em R$_teste.",
        )
        generator = DocxReportGenerator(config)
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            logs_dir = root / "logs"
            logs_dir.mkdir(parents=True)
            (logs_dir / "run.log").write_text("linha de log", encoding="utf-8")
            result = build_inventory_result(
                root,
                logs_dir,
                source_name="video_50%_&_relatorio_#1.mp4",
            )

            artifacts = generator.generate(result)

            full_text = _docx_full_text(artifacts.docx_path)
            self.assertIn("Nota com 100% de cadeia & custos #1 em R$_teste.", full_text)


if __name__ == "__main__":
    unittest.main()
